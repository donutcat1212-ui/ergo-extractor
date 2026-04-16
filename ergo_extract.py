
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Извлечение данных из выписных эрготерапевтических осмотров (.docx)
и формирование отдельной итоговой Excel-таблицы.

Что делает:
1) Ищет docx только в папках выписных/повторных осмотров.
2) Читает Ф.И.О. и дату/время документа из текста Word.
3) Извлекает значения по шкалам и доменам МКФ из таблиц документа.
4) Если по одному пациенту и одной дате найдено несколько файлов, выбирает самый поздний по дате/времени.
5) Создает отдельный Excel-файл с листами:
   - cases_wide   : одна строка = один случай
   - values_long  : длинный формат по каждому показателю
   - files_raw    : все обработанные файлы
   - warnings     : предупреждения и спорные случаи

Зависимости:
    pip install python-docx openpyxl
"""

from __future__ import annotations

import argparse
import datetime as dt
import os
import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


# -----------------------------
# БАЗОВЫЕ НАСТРОЙКИ ПО УМОЛЧАНИЮ
# -----------------------------
DEFAULT_ROOT_DIR = r"\\fccps.local\dfs\ОМР ЦНС1\Эрготерапия\Эрготерапия 23-25г"
DEFAULT_TEMPLATE_XLSX = r"ЭРГО(1).xlsx"   # опционально, можно оставить пустым
DEFAULT_OUTPUT_XLSX = r"ergotherapy_extract_result.xlsx"


# ---------------------------------
# СПИСОК ПОЛЕЙ ПО УМОЛЧАНИЮ (fallback)
# ---------------------------------
DEFAULT_FIELD_DEFINITIONS = [
    ("Шкала SULCS", "", "paired"),
    ("Шкала SULCS", "противоположная рука", "paired"),
    ("Шкала FIM", "", "paired"),
    ("Шкала EQ-5D", "", "paired"),
    ("ОЦЕНКА COPM", "", "paired"),
    ("d445", "Использование кисти руки", "paired"),
    ("d445", "Использование кисти руки (противоположная рука)", "paired"),
    ("d4458", "Использование кисти и руки, другое уточненное", "paired"),
    ("d540", "Надевание одежды", "paired"),
    ("d6300", "Приготовление простых блюд", "paired"),
    ("е1151", "Вспомогательные изделия и технологии для личного повседневного пользования", "paired"),
    ("е155", "Дизайн, характер проектирования, строительства и обустройства зданий частного использования", "paired"),
    ("е150", "Дизайн, характер проектирования, строительства и обустройства зданий для общественного использования", "paired"),
    ("d4408", "Использование точных движений кисти, другое уточненное (удержание адаптационной ручки при письме правой рукой)", "paired"),
    ("d4301", "Перенос кистями рук", "paired"),
    ("d145", "Усвоение навыков письма (четкое написание букв правой рукой)", "paired"),
    ("d6201", "Приготовление сложных блюд", "paired"),
    ("d550", "Прием пищи (использование столовых приборов правой рукой)", "paired"),
    ("d2302", "Исполнение повседневного распорядка", "paired"),
    ("d5208", "Уход за частями тела, другой уточненный (паретичная верхняя конечность)", "paired"),
    ("d4108", "Изменение позы тела уточненное", "paired"),
    ("d4208", "Перемещение тела другое, уточненное", "paired"),
    ("d5202", "Уход за волосами", "paired"),
    ("d510", "Мытье", "paired"),
    ("d530", "Физиологические отправления", "paired"),
    ("d129", "Целенаправленное использование органов чувств, другое уточненное и не уточненное", "paired"),
    ("d4158", "Поддержание положения тела", "paired"),
    ("d460", "Передвижение в различных местах", "paired"),
    ("d179", "Применение знаний, другое уточненное", "paired"),
    ("норма", "", "activity"),
]


RU_STOPWORDS = {
    "и", "или", "для", "при", "другое", "другой", "уточненное", "уточненный",
    "характер", "дизайн", "обустройства", "строительства", "использование",
    "личного", "повседневного", "уход", "за", "в", "на", "по", "от", "до",
    "рука", "руки", "рукой", "конечность", "правая", "левой", "правой",
    "верхняя", "частного", "общественного", "положения", "тела", "другое",
    "другое", "другое", "уточненное", "не", "поддержание", "изменение",
    "перемещение", "точных", "движений", "кисти", "органов", "чувств",
}

CYR_TO_LAT = str.maketrans({
    "е": "e", "Е": "E",
    "а": "a", "А": "A",
    "о": "o", "О": "O",
    "р": "p", "Р": "P",
    "с": "c", "С": "C",
    "х": "x", "Х": "X",
    "у": "y", "У": "Y",
    "к": "k", "К": "K",
    "м": "m", "М": "M",
    "т": "t", "Т": "T",
})


@dataclass
class FieldSpec:
    code: str
    description: str
    kind: str = "paired"   # paired / activity
    display: str = ""
    field_id: str = ""
    code_norm: str = ""
    display_norm: str = ""
    desc_norm: str = ""
    desc_tokens: Tuple[str, ...] = field(default_factory=tuple)

    def __post_init__(self) -> None:
        code = clean_text(self.code)
        description = clean_text(self.description)
        self.code = code
        self.description = description
        self.display = self.display or make_display_name(code, description, self.kind)
        self.field_id = self.field_id or build_field_id(code, description, self.kind)
        self.code_norm = normalize_code_text(code)
        self.display_norm = normalize_match(self.display)
        self.desc_norm = normalize_match(description)
        self.desc_tokens = tuple(tokenize_desc(self.description))


@dataclass
class MetricCandidate:
    field_id: str
    display: str
    value_in: Optional[str]
    value_out: Optional[str]
    score: int
    matched_by: str
    source: str
    row_text: str


@dataclass
class ParsedDocument:
    file_path: str
    year: str
    patient: str
    surname: str
    patient_norm: str
    exam_datetime: Optional[dt.datetime]
    exam_date: Optional[dt.date]
    activity: str
    selected_warning: str
    values: Dict[str, Tuple[Optional[str], Optional[str]]]
    metric_debug: Dict[str, MetricCandidate]
    warnings: List[str]


# -----------------------------
# НОРМАЛИЗАЦИЯ
# -----------------------------
def clean_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", " ", text)
    return text.strip()


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", clean_text(text)).strip()


def normalize_match(text: str) -> str:
    text = clean_text(text).lower().replace("ё", "е")
    text = text.replace("—", "-").replace("–", "-")
    text = re.sub(r"[/\\|;:,()\[\]{}]+", " ", text)
    text = re.sub(r"[^0-9a-zа-я\-\. ]+", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def normalize_code_text(text: str) -> str:
    text = normalize_match(text)
    return text.translate(CYR_TO_LAT)


def tokenize_desc(text: str) -> List[str]:
    raw = re.findall(r"[a-zа-я0-9]+", normalize_match(text))
    out = []
    for token in raw:
        if len(token) <= 2:
            continue
        if token in RU_STOPWORDS:
            continue
        out.append(token)
    return out


def build_field_id(code: str, description: str, kind: str) -> str:
    parts = [normalize_match(code)]
    if description:
        parts.append(normalize_match(description))
    if kind == "activity":
        parts = ["activity"]
    compact = "__".join([p for p in parts if p])
    compact = compact.replace(" ", "_")
    return compact or "field"


def make_display_name(code: str, description: str, kind: str) -> str:
    if kind == "activity":
        return "Характер занятий"
    if description:
        if code and description and normalize_match(code) not in normalize_match(description):
            return f"{code} | {description}"
        return description
    return code or "Поле"


def surname_from_fio(fio: str) -> str:
    fio = normalize_spaces(fio)
    if not fio:
        return ""
    parts = fio.split()
    return parts[0] if parts else fio


def normalize_patient_name(text: str) -> str:
    text = normalize_match(text)
    text = re.sub(r"\b(пациент|больной|ф и о|фио)\b", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def patient_key_from_name_date(name: str, date_obj: Optional[dt.date], year: str) -> str:
    date_part = date_obj.isoformat() if date_obj else "unknown_date"
    return f"{year}|{normalize_patient_name(name)}|{date_part}"


# -----------------------------
# ЗАГРУЗКА ПОЛЕЙ
# -----------------------------
def load_fields_from_template(template_xlsx: str) -> List[FieldSpec]:
    fields: List[FieldSpec] = []
    if not template_xlsx:
        return fields
    if not os.path.exists(template_xlsx):
        return fields

    try:
        wb = load_workbook(template_xlsx, data_only=True)
    except Exception:
        return fields

    best_ws = None
    best_score = -1
    for ws in wb.worksheets:
        score = 0
        for c in range(3, ws.max_column + 1, 2):
            if clean_text(ws.cell(3, c).value) or clean_text(ws.cell(4, c).value):
                score += 1
        if score > best_score:
            best_score = score
            best_ws = ws

    if best_ws is None:
        return fields

    for c in range(3, best_ws.max_column + 1, 2):
        code = clean_text(best_ws.cell(3, c).value)
        desc = clean_text(best_ws.cell(4, c).value)
        if not code and not desc:
            continue
        fields.append(FieldSpec(code=code, description=desc, kind="paired"))

    return fields


def get_target_fields(template_xlsx: str) -> List[FieldSpec]:
    fields_by_id: Dict[str, FieldSpec] = {}

    for code, desc, kind in DEFAULT_FIELD_DEFINITIONS:
        field = FieldSpec(code=code, description=desc, kind=kind)
        fields_by_id[field.field_id] = field

    for field in load_fields_from_template(template_xlsx):
        if field.field_id not in fields_by_id:
            fields_by_id[field.field_id] = field

    if "activity" not in fields_by_id:
        activity_field = FieldSpec(code="норма", description="", kind="activity")
        fields_by_id[activity_field.field_id] = activity_field

    ordered = list(fields_by_id.values())

    # activity в конец
    paired = [f for f in ordered if f.kind != "activity"]
    activity = [f for f in ordered if f.kind == "activity"]
    return paired + activity


# -----------------------------
# DOCX ЧТЕНИЕ
# -----------------------------
def extract_docx_content(doc_path: str) -> Tuple[List[str], List[List[List[str]]], str]:
    doc = Document(doc_path)

    paragraphs = [clean_text(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]

    tables: List[List[List[str]]] = []
    table_texts: List[str] = []

    for table in doc.tables:
        table_rows: List[List[str]] = []
        for row in table.rows:
            row_cells = [clean_text(cell.text) for cell in row.cells]
            table_rows.append(row_cells)
            table_texts.extend([c for c in row_cells if c])
        tables.append(table_rows)

    combined_text = "\n".join(paragraphs + table_texts)
    return paragraphs, tables, combined_text


# -----------------------------
# ПАРСИНГ РЕКВИЗИТОВ ДОКУМЕНТА
# -----------------------------
PATIENT_PATTERNS = [
    re.compile(r"Ф\.?\s*И\.?\s*О\.?\s*:\s*(.+?)(?:\n|$)", re.IGNORECASE),
    re.compile(r"ФИО\s*:\s*(.+?)(?:\n|$)", re.IGNORECASE),
]

DATETIME_PATTERNS = [
    re.compile(r"дата\s*/\s*время\s*[:\-]?\s*([0-3]?\d\.[01]?\d\.\d{4}\s*/\s*[0-2]?\d:[0-5]\d)", re.IGNORECASE),
    re.compile(r"дата\s*время\s*[:\-]?\s*([0-3]?\d\.[01]?\d\.\d{4}\s*/\s*[0-2]?\d:[0-5]\d)", re.IGNORECASE),
]


def parse_patient_from_text(text: str, filename: str) -> str:
    for pattern in PATIENT_PATTERNS:
        m = pattern.search(text)
        if m:
            patient = clean_text(m.group(1))
            patient = re.split(r"\b(?:дата\s*/?\s*время|возраст|палата|адрес|жалобы)\b", patient, flags=re.IGNORECASE)[0]
            patient = patient.strip(" -:;.,")
            if patient:
                return normalize_spaces(patient)

    stem = Path(filename).stem
    stem = re.sub(r"^\d+\s*", "", stem)
    stem = re.sub(r"\b(выписка|консультация|эрго|эт|эрготерапия)\b", " ", stem, flags=re.IGNORECASE)
    stem = re.sub(r"[_\-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip(" ._-")
    return normalize_spaces(stem)


def parse_datetime_from_text(text: str) -> Optional[dt.datetime]:
    for pattern in DATETIME_PATTERNS:
        m = pattern.search(text)
        if not m:
            continue
        raw = m.group(1)
        raw = raw.replace(" ", "")
        try:
            return dt.datetime.strptime(raw, "%d.%m.%Y/%H:%M")
        except ValueError:
            continue
    return None


def detect_year_from_path(path: str) -> str:
    normalized = path.replace("\\", "/")
    m = re.search(r"(20\d{2})", normalized)
    return m.group(1) if m else ""


def is_discharge_dir_name(name: str) -> bool:
    n = normalize_match(name)
    return ("выпис" in n) or ("повтор" in n)


def iter_discharge_docx_files(root_dir: str) -> Iterable[str]:
    root = Path(root_dir)
    if not root.exists():
        return []

    all_paths: List[str] = []
    for current_root, dirnames, filenames in os.walk(root_dir):
        current_name = os.path.basename(current_root)
        if is_discharge_dir_name(current_name):
            for fn in filenames:
                if fn.lower().endswith(".docx") and not fn.startswith("~$"):
                    all_paths.append(os.path.join(current_root, fn))
    return sorted(all_paths)


# -----------------------------
# ПАРСИНГ ЗНАЧЕНИЙ
# -----------------------------
def find_header_columns(table_rows: List[List[str]]) -> Tuple[Optional[int], Optional[int]]:
    in_keywords = ["поступ", "первич", "при приеме", "при приёме", "при поступ"]
    out_keywords = ["выписк", "повтор", "при выписк", "при выпис", "на выписк"]

    in_idx = None
    out_idx = None

    for row in table_rows[:4]:
        for idx, cell in enumerate(row):
            n = normalize_match(cell)
            if not n:
                continue
            if in_idx is None and any(k in n for k in in_keywords):
                in_idx = idx
            if out_idx is None and any(k in n for k in out_keywords):
                out_idx = idx

    return in_idx, out_idx


COPM_PAIR_RE = re.compile(r"(?<!\d)(\d+(?:[.,]\d+)?)\s*[-–—]\s*(\d+(?:[.,]\d+)?)(?!\d)")
NUMBER_RE = re.compile(r"(?<![\d/])\d+(?:[.,]\d+)?(?![:\d])")


def normalize_number_text(value: str) -> str:
    value = clean_text(value)
    value = value.replace(",", ".")
    value = re.sub(r"\s+", "", value)
    return value


def parse_value_text(cell_text: str) -> Optional[str]:
    text = clean_text(cell_text)
    if not text:
        return None

    pair = COPM_PAIR_RE.search(text)
    if pair:
        left = normalize_number_text(pair.group(1))
        right = normalize_number_text(pair.group(2))
        return f"{left}-{right}"

    if text in {"-", "—", "–"}:
        return None

    # иногда в ячейке есть фраза вида "44" / "2,5"
    nums = NUMBER_RE.findall(text)
    if len(nums) == 1:
        return normalize_number_text(nums[0])

    return None


def extract_last_two_scalar_values(row_cells: Sequence[str]) -> Tuple[Optional[str], Optional[str]]:
    scalar_cells: List[str] = []
    for cell in row_cells:
        parsed = parse_value_text(cell)
        if parsed is not None:
            scalar_cells.append(parsed)

    if not scalar_cells:
        joined = " | ".join([clean_text(c) for c in row_cells if clean_text(c)])
        pair = COPM_PAIR_RE.search(joined)
        if pair:
            return normalize_number_text(pair.group(1)), normalize_number_text(pair.group(2))

        nums = [normalize_number_text(x) for x in NUMBER_RE.findall(joined)]
        if len(nums) >= 2:
            return nums[-2], nums[-1]
        if len(nums) == 1:
            return nums[0], None
        return None, None

    if len(scalar_cells) >= 2:
        return scalar_cells[-2], scalar_cells[-1]
    return scalar_cells[0], None


def row_match_score(field: FieldSpec, row_cells: Sequence[str]) -> Tuple[int, str]:
    row_joined = " | ".join([clean_text(c) for c in row_cells if clean_text(c)])
    row_norm = normalize_match(row_joined)
    row_code_norm = normalize_code_text(row_joined)

    if not row_norm:
        return 0, ""

    score = 0
    matched_by = []

    code_hit = False
    if field.code_norm and field.code_norm in row_code_norm:
        score += 5
        matched_by.append("code")
        code_hit = True

    if field.desc_norm and field.desc_norm in row_norm:
        score += 6
        matched_by.append("desc_exact")
    elif field.desc_tokens:
        overlap = sum(1 for t in field.desc_tokens if t in row_norm)
        if overlap >= 2:
            score += min(5, overlap + 1)
            matched_by.append(f"desc_tokens:{overlap}")

    if field.display_norm and field.display_norm in row_norm:
        score += 4
        matched_by.append("display")

    # защита от ложных совпадений для повторяющихся кодов типа d445
    if field.description and code_hit and field.desc_tokens:
        overlap = sum(1 for t in field.desc_tokens if t in row_norm)
        if overlap == 0 and field.desc_norm not in row_norm:
            score -= 3

    # шкалы
    if field.code and "шкала" in normalize_match(field.code) and normalize_match(field.code) in row_norm:
        score += 2
        matched_by.append("scale")

    return score, ",".join(matched_by)


def extract_pair_from_row(row_cells: Sequence[str], in_idx: Optional[int], out_idx: Optional[int]) -> Tuple[Optional[str], Optional[str], str]:
    value_in = None
    value_out = None
    mode = "fallback"

    if in_idx is not None and in_idx < len(row_cells):
        value_in = parse_value_text(row_cells[in_idx])
    if out_idx is not None and out_idx < len(row_cells):
        value_out = parse_value_text(row_cells[out_idx])

    if value_in is not None or value_out is not None:
        return value_in, value_out, "header_columns"

    value_in, value_out = extract_last_two_scalar_values(row_cells)
    return value_in, value_out, mode


def detect_activity(text: str) -> Tuple[str, str]:
    text_norm = normalize_match(text)
    found = []

    if "группов" in text_norm:
        found.append("групповые")
    if "индивидуаль" in text_norm:
        found.append("индивидуальные")
    if "консультац" in text_norm:
        found.append("консультация")

    found = list(dict.fromkeys(found))

    if not found:
        return "", ""
    if len(found) == 1:
        return found[0], ""

    # неоднозначность лучше не угадывать
    return "", f"Найдено несколько типов активности: {', '.join(found)}"


def parse_metrics_from_doc(
    paragraphs: List[str],
    tables: List[List[List[str]]],
    full_text: str,
    fields: List[FieldSpec],
) -> Tuple[Dict[str, Tuple[Optional[str], Optional[str]]], Dict[str, MetricCandidate], List[str], str]:
    warnings: List[str] = []
    activity, activity_warning = detect_activity(full_text)
    if activity_warning:
        warnings.append(activity_warning)

    paired_fields = [f for f in fields if f.kind == "paired"]

    best_candidates: Dict[str, MetricCandidate] = {}

    # сначала таблицы
    for t_index, table_rows in enumerate(tables):
        in_idx, out_idx = find_header_columns(table_rows)

        for r_index, row_cells in enumerate(table_rows):
            row_text = " | ".join([clean_text(c) for c in row_cells if clean_text(c)])
            if not row_text:
                continue

            for field in paired_fields:
                score, matched_by = row_match_score(field, row_cells)
                if score < 5:
                    continue

                value_in, value_out, mode = extract_pair_from_row(row_cells, in_idx, out_idx)
                if value_in is None and value_out is None:
                    continue

                candidate = MetricCandidate(
                    field_id=field.field_id,
                    display=field.display,
                    value_in=value_in,
                    value_out=value_out,
                    score=score + (2 if mode == "header_columns" else 0),
                    matched_by=f"table:{matched_by};{mode}",
                    source=f"table_{t_index + 1}_row_{r_index + 1}",
                    row_text=row_text,
                )

                existing = best_candidates.get(field.field_id)
                if existing is None or candidate.score > existing.score:
                    best_candidates[field.field_id] = candidate

    # fallback: поиск по строкам общего текста
    lines = [clean_text(line) for line in full_text.splitlines()]
    lines = [line for line in lines if line]

    for line_index, line in enumerate(lines):
        fake_row = [line]
        for field in paired_fields:
            if field.field_id in best_candidates:
                continue
            score, matched_by = row_match_score(field, fake_row)
            if score < 5:
                continue

            pair = COPM_PAIR_RE.search(line)
            if pair:
                value_in = normalize_number_text(pair.group(1))
                value_out = normalize_number_text(pair.group(2))
            else:
                nums = [normalize_number_text(x) for x in NUMBER_RE.findall(line)]
                if len(nums) >= 2:
                    value_in, value_out = nums[-2], nums[-1]
                elif len(nums) == 1:
                    value_in, value_out = nums[0], None
                else:
                    continue

            best_candidates[field.field_id] = MetricCandidate(
                field_id=field.field_id,
                display=field.display,
                value_in=value_in,
                value_out=value_out,
                score=score,
                matched_by=f"line:{matched_by}",
                source=f"line_{line_index + 1}",
                row_text=line,
            )

    values: Dict[str, Tuple[Optional[str], Optional[str]]] = {}
    for field in paired_fields:
        cand = best_candidates.get(field.field_id)
        if cand:
            values[field.field_id] = (cand.value_in, cand.value_out)
        else:
            values[field.field_id] = ("", "")

    return values, best_candidates, warnings, activity


# -----------------------------
# ОБРАБОТКА ФАЙЛА
# -----------------------------
def parse_single_doc(doc_path: str, fields: List[FieldSpec]) -> ParsedDocument:
    warnings: List[str] = []
    year = detect_year_from_path(doc_path)

    paragraphs, tables, full_text = extract_docx_content(doc_path)
    patient = parse_patient_from_text(full_text, os.path.basename(doc_path))
    patient_norm = normalize_patient_name(patient)
    surname = surname_from_fio(patient)
    exam_datetime = parse_datetime_from_text(full_text)

    if not patient:
        warnings.append("Не удалось извлечь Ф.И.О.")
    if exam_datetime is None:
        warnings.append("Не удалось извлечь дату/время")

    values, metric_debug, metric_warnings, activity = parse_metrics_from_doc(
        paragraphs=paragraphs,
        tables=tables,
        full_text=full_text,
        fields=fields,
    )
    warnings.extend(metric_warnings)

    if not any((v_in or v_out) for v_in, v_out in values.values()):
        warnings.append("Не удалось извлечь ни одного показателя")

    return ParsedDocument(
        file_path=doc_path,
        year=year,
        patient=patient,
        surname=surname,
        patient_norm=patient_norm,
        exam_datetime=exam_datetime,
        exam_date=exam_datetime.date() if exam_datetime else None,
        activity=activity,
        selected_warning="",
        values=values,
        metric_debug=metric_debug,
        warnings=warnings,
    )


# -----------------------------
# ДЕДУПЛИКАЦИЯ
# -----------------------------
def deduplicate_documents(docs: List[ParsedDocument]) -> Tuple[List[ParsedDocument], List[Tuple[str, str, str]]]:
    groups: Dict[str, List[ParsedDocument]] = defaultdict(list)
    for doc in docs:
        key = patient_key_from_name_date(doc.patient or doc.surname, doc.exam_date, doc.year)
        groups[key].append(doc)

    selected_docs: List[ParsedDocument] = []
    duplicates_log: List[Tuple[str, str, str]] = []

    for key, group in groups.items():
        group_sorted = sorted(
            group,
            key=lambda d: (
                d.exam_datetime or dt.datetime.min,
                len([1 for pair in d.values.values() if pair != ("", "")]),
            ),
            reverse=True,
        )
        winner = group_sorted[0]
        if len(group_sorted) > 1:
            winner.selected_warning = f"Найдено файлов в группе: {len(group_sorted)}. Выбран самый поздний."
            for loser in group_sorted[1:]:
                duplicates_log.append((key, winner.file_path, loser.file_path))
        selected_docs.append(winner)

    selected_docs.sort(key=lambda d: ((d.year or ""), d.exam_date or dt.date.min, d.patient_norm))
    return selected_docs, duplicates_log


# -----------------------------
# ЭКСПОРТ В EXCEL
# -----------------------------
def autosize_columns(ws) -> None:
    widths: Dict[int, int] = {}
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            width = min(max(len(str(cell.value)) + 2, 10), 60)
            widths[cell.column] = max(widths.get(cell.column, 0), width)
    for col_idx, width in widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width


def style_header(ws, header_row: int = 1) -> None:
    fill = PatternFill("solid", fgColor="1F4E78")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[header_row]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"


def export_result_xlsx(
    output_path: str,
    selected_docs: List[ParsedDocument],
    all_docs: List[ParsedDocument],
    fields: List[FieldSpec],
    duplicates_log: List[Tuple[str, str, str]],
) -> None:
    wb = Workbook()

    # 1) wide
    ws = wb.active
    ws.title = "cases_wide"

    paired_fields = [f for f in fields if f.kind == "paired"]

    headers = [
        "Год",
        "Дата осмотра",
        "Дата/время документа",
        "Пациент",
        "Фамилия",
        "Нормализованный пациент",
        "Ключ сопоставления",
        "Характер занятий",
        "Источник файла",
        "Предупреждение выбора",
        "Число извлеченных показателей",
    ]
    for field in paired_fields:
        headers.append(f"{field.display} | поступление")
        headers.append(f"{field.display} | выписка")

    ws.append(headers)

    for doc in selected_docs:
        row = [
            doc.year,
            doc.exam_date.isoformat() if doc.exam_date else "",
            doc.exam_datetime.strftime("%d.%m.%Y %H:%M") if doc.exam_datetime else "",
            doc.patient,
            doc.surname,
            doc.patient_norm,
            patient_key_from_name_date(doc.patient or doc.surname, doc.exam_date, doc.year),
            doc.activity,
            doc.file_path,
            doc.selected_warning,
            sum(1 for value_in, value_out in doc.values.values() if value_in or value_out),
        ]
        for field in paired_fields:
            value_in, value_out = doc.values.get(field.field_id, ("", ""))
            row.extend([value_in or "", value_out or ""])
        ws.append(row)

    style_header(ws)
    autosize_columns(ws)

    # 2) long
    ws_long = wb.create_sheet("values_long")
    ws_long.append([
        "Год", "Дата осмотра", "Пациент", "Фамилия",
        "Поле", "Значение при поступлении", "Значение при выписке",
        "Источник файла", "Источник совпадения", "Оценка совпадения", "Строка-источник"
    ])

    for doc in selected_docs:
        for field in paired_fields:
            cand = doc.metric_debug.get(field.field_id)
            value_in, value_out = doc.values.get(field.field_id, ("", ""))
            if not value_in and not value_out:
                continue
            ws_long.append([
                doc.year,
                doc.exam_date.isoformat() if doc.exam_date else "",
                doc.patient,
                doc.surname,
                field.display,
                value_in or "",
                value_out or "",
                doc.file_path,
                cand.matched_by if cand else "",
                cand.score if cand else "",
                cand.row_text if cand else "",
            ])

    style_header(ws_long)
    autosize_columns(ws_long)

    # 3) all files
    ws_files = wb.create_sheet("files_raw")
    ws_files.append([
        "Год", "Дата/время документа", "Дата осмотра", "Пациент", "Фамилия",
        "Нормализованный пациент", "Источник файла", "Число извлеченных показателей",
        "Характер занятий", "Предупреждения"
    ])

    for doc in all_docs:
        ws_files.append([
            doc.year,
            doc.exam_datetime.strftime("%d.%m.%Y %H:%M") if doc.exam_datetime else "",
            doc.exam_date.isoformat() if doc.exam_date else "",
            doc.patient,
            doc.surname,
            doc.patient_norm,
            doc.file_path,
            sum(1 for value_in, value_out in doc.values.values() if value_in or value_out),
            doc.activity,
            " | ".join(doc.warnings),
        ])

    style_header(ws_files)
    autosize_columns(ws_files)

    # 4) warnings
    ws_warn = wb.create_sheet("warnings")
    ws_warn.append(["Тип", "Год", "Пациент", "Дата осмотра", "Файл", "Сообщение"])

    for doc in all_docs:
        for warning in doc.warnings:
            ws_warn.append([
                "document_warning",
                doc.year,
                doc.patient,
                doc.exam_date.isoformat() if doc.exam_date else "",
                doc.file_path,
                warning,
            ])

    for key, winner, loser in duplicates_log:
        ws_warn.append([
            "duplicate_group",
            key.split("|", 1)[0],
            "",
            "",
            loser,
            f"Дубликат группы. Оставлен файл: {winner}",
        ])

    style_header(ws_warn)
    autosize_columns(ws_warn)

    wb.save(output_path)


# -----------------------------
# MAIN
# -----------------------------
def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Извлечение шкал и доменов из выписных DOCX и формирование отдельной Excel-таблицы."
    )
    parser.add_argument("--root", default=DEFAULT_ROOT_DIR, help="Корневая папка с годами и выписными осмотрами")
    parser.add_argument("--template", default=DEFAULT_TEMPLATE_XLSX, help="Excel-шаблон для ориентировочных заголовков (опционально)")
    parser.add_argument("--output", default=DEFAULT_OUTPUT_XLSX, help="Путь к итоговому Excel")
    return parser


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()

    if not os.path.exists(args.root):
        print(f"[ОШИБКА] Корневая папка не найдена: {args.root}")
        return 2

    fields = get_target_fields(args.template)
    if not fields:
        print("[ОШИБКА] Не удалось сформировать список полей.")
        return 2

    files = list(iter_discharge_docx_files(args.root))
    if not files:
        print("[ОШИБКА] Не найдено ни одного .docx в папках выписных осмотров.")
        return 2

    print(f"[INFO] Найдено файлов: {len(files)}")

    parsed_docs: List[ParsedDocument] = []
    for index, file_path in enumerate(files, start=1):
        try:
            doc = parse_single_doc(file_path, fields)
            parsed_docs.append(doc)
            if index % 25 == 0 or index == len(files):
                print(f"[INFO] Обработано: {index}/{len(files)}")
        except Exception as exc:
            parsed_docs.append(ParsedDocument(
                file_path=file_path,
                year=detect_year_from_path(file_path),
                patient="",
                surname="",
                patient_norm="",
                exam_datetime=None,
                exam_date=None,
                activity="",
                selected_warning="",
                values={field.field_id: ("", "") for field in fields if field.kind == "paired"},
                metric_debug={},
                warnings=[f"Критическая ошибка разбора файла: {exc}"],
            ))

    selected_docs, duplicates_log = deduplicate_documents(parsed_docs)
    export_result_xlsx(args.output, selected_docs, parsed_docs, fields, duplicates_log)

    print(f"[OK] Готово. Итоговый файл: {args.output}")
    print(f"[OK] Уникальных случаев: {len(selected_docs)}")
    print(f"[OK] Всего обработанных файлов: {len(parsed_docs)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
