
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
ergo_extract_v2.py

Улучшенный извлекатель данных из первичных и выписных осмотров/консультаций эрготерапии.
Создает отдельный Excel-файл и не трогает исходную таблицу.

Главные отличия от первой версии:
- читает и первичные, и выписные документы;
- корректно понимает дату/время вида 05.06.2025/16.02 и 05.06.2025/16:02;
- для SULCS/FIM/EQ-5D/COPM приоритетно берет итоговую сравнительную таблицу с датами;
- не затирает парные значения одиночным числом из описательной части;
- умеет оставлять только "поступление" для первичных консультаций без выписки;
- характер занятий определяет по смысловой финальной части документа, а не по заголовку "Консультация ...";
- добавляет первичные осмотры только там, где нет соответствующей выписки.

Запуск:
    ergo_extract_v2.exe --root ".\\Эрготерапия 23-25г"
или
    python ergo_extract_v2.py --root ".\\Эрготерапия 23-25г"
"""

from __future__ import annotations

import argparse
import datetime as dt
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

DEFAULT_ROOT_DIR = r"\\fccps.local\dfs\ОМР ЦНС1\Эрготерапия\Эрготерапия 23-25г"
DEFAULT_OUTPUT_XLSX = r"ergotherapy_extract_result_v2.xlsx"

FIELD_DEFINITIONS = [
    ("Шкала SULCS", "", "paired"),
    ("Шкала FIM", "", "paired"),
    ("Шкала EQ-5D", "", "paired"),
    ("ОЦЕНКА COPM", "", "paired"),
    ("d445", "Использование кисти руки", "paired"),
    ("d4458", "Использование кисти и руки, другое уточненное", "paired"),
    ("d540", "Надевание одежды", "paired"),
    ("d6300", "Приготовление простых блюд", "paired"),
    ("е1151", "Вспомогательные изделия и технологии для личного повседневного пользования", "paired"),
    ("е155", "Дизайн, характер проектирования, строительства и обустройства зданий частного использования", "paired"),
    ("е150", "Дизайн, характер проектирования, строительства и обустройства зданий для общественного использования", "paired"),
    ("d4408", "Использование точных движений кисти, другое уточненное (удержание адаптационной ручки при письме правой рукой)", "paired"),
    ("d4301", "Перенос кистями рук", "paired"),
    ("d145", "Усвоение навыков письма (четкое написание букв правой рукой)", "paired"),
    ("d6201", "Приготовление сложных блюд", "paired"),
    ("d550", "Прием пищи (использование столовых приборов правой рукой)", "paired"),
    ("d2302", "Исполнение повседневного распорядка", "paired"),
    ("d5208", "Уход за частями тела, другой уточненный (паретичная верхняя конечность)", "paired"),
    ("d4108", "Изменение позы тела уточненное", "paired"),
    ("d4208", "Перемещение тела другое, уточненное", "paired"),
    ("d5202", "Уход за волосами", "paired"),
    ("d510", "Мытье", "paired"),
    ("d530", "Физиологические отправления", "paired"),
    ("d129", "Целенаправленное использование органов чувств, другое уточненное и не уточненное", "paired"),
    ("d4158", "Поддержание положения тела", "paired"),
    ("d460", "Передвижение в различных местах", "paired"),
    ("d179", "Применение знаний, другое уточненное", "paired"),
    ("норма", "", "activity"),
]

RU_STOPWORDS = {
    "и", "или", "для", "при", "другое", "другой", "уточненное", "уточненный",
    "характер", "дизайн", "обустройства", "строительства", "использование",
    "личного", "повседневного", "уход", "за", "в", "на", "по", "от", "до",
    "рука", "руки", "рукой", "конечность", "верхняя", "верхних", "верхних",
    "нижняя", "нижних", "правая", "правой", "левой", "кисти", "тела", "другое",
    "уточненный", "уточненное", "функциональной", "независимости"
}

DATE_RE = re.compile(r'(?<!\d)(\d{2}\.\d{2}\.\d{4})(?!\d)')
DATETIME_PATTERNS = [
    re.compile(r'дата/время\s*[:\-]?\s*(\d{2}\.\d{2}\.\d{4}\s*/\s*\d{2}[:.]\d{2})', re.I),
    re.compile(r'дата\s*[:\-]?\s*(\d{2}\.\d{2}\.\d{4}\s*/\s*\d{2}[:.]\d{2})', re.I),
]
COPM_PAIR_RE = re.compile(r'(?<!\d)(\d+(?:[.,]\d+)?)\s*[-–—]\s*(\d+(?:[.,]\d+)?)(?!\d)')
NUMBER_RE = re.compile(r'(?<![\d/])\d+(?:[.,]\d+)?(?![:\d])')
CODE_RE = re.compile(r'^[deе]\d{3,4}$', re.I)


def clean_text(text: str) -> str:
    text = text or ""
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s*\n\s*", "\n", text)
    return text.strip()


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", clean_text(text)).strip()


def normalize_match(text: str) -> str:
    text = normalize_spaces(text).lower()
    text = text.replace("ё", "е")
    text = text.replace("№", " ")
    return text


def normalize_code_text(text: str) -> str:
    text = normalize_match(text)
    # латинская e и кириллическая е -> кириллическая е
    text = text.replace("e", "е")
    text = text.replace("d ", "d")
    text = re.sub(r'([dе])\s+(\d)', r'\1\2', text)
    return text


def normalize_number_text(text: str) -> str:
    text = clean_text(text)
    text = text.replace(",", ".")
    return text


def parse_value_text(cell_text: str) -> Optional[str]:
    text = normalize_spaces(cell_text)
    if not text or text in {"-", "—", "–"}:
        return None

    pair = COPM_PAIR_RE.search(text)
    if pair:
        left = normalize_number_text(pair.group(1))
        right = normalize_number_text(pair.group(2))
        return f"{left}-{right}"

    nums = NUMBER_RE.findall(text)
    if len(nums) == 1:
        return normalize_number_text(nums[0])
    return None


def tokenize_desc(text: str) -> Tuple[str, ...]:
    text = normalize_match(text)
    tokens = re.findall(r"[a-zA-Zа-яА-ЯеЕёЁ0-9]+", text)
    out = []
    for t in tokens:
        if t in RU_STOPWORDS:
            continue
        if len(t) <= 1:
            continue
        out.append(t)
    return tuple(dict.fromkeys(out))


def build_field_id(code: str, desc: str, kind: str) -> str:
    return f"{normalize_code_text(code)}|{normalize_match(desc)}|{kind}"


def make_display_name(code: str, desc: str, kind: str) -> str:
    if kind == "activity":
        return code
    return f"{code} | {desc}" if desc else code


def surname_from_fio(fio: str) -> str:
    fio = normalize_spaces(fio)
    return fio.split()[0] if fio else ""


def normalize_patient_name(name: str) -> str:
    return normalize_match(name)


def patient_key_from_name_date(year: str, name: str, date_obj: Optional[dt.date]) -> str:
    d = date_obj.isoformat() if date_obj else "unknown_date"
    return f"{year}|{normalize_patient_name(name)}|{d}"


def canonical_code(raw: str) -> str:
    code = normalize_code_text(raw)
    # убрать лишний ноль на конце для d5400 -> d540 и d5100 -> d510
    if re.fullmatch(r'[dе]\d{4}', code):
        if code[-1] == '0':
            code = code[:-1]
    return code


@dataclass
class FieldSpec:
    code: str
    description: str
    kind: str = "paired"
    display: str = ""
    field_id: str = ""
    code_norm: str = ""
    desc_norm: str = ""
    desc_tokens: Tuple[str, ...] = field(default_factory=tuple)

    def __post_init__(self):
        self.code = clean_text(self.code)
        self.description = clean_text(self.description)
        self.display = self.display or make_display_name(self.code, self.description, self.kind)
        self.code_norm = canonical_code(self.code)
        self.desc_norm = normalize_match(self.description)
        self.desc_tokens = tokenize_desc(self.description)
        self.field_id = build_field_id(self.code_norm, self.description, self.kind)


@dataclass
class ParsedDoc:
    path: str
    year: str
    doc_kind: str  # primary/discharge
    patient: str
    surname: str
    patient_norm: str
    document_dt: Optional[dt.datetime]
    exam_date: Optional[dt.date]
    primary_eval_date: Optional[dt.date]
    repeat_eval_date: Optional[dt.date]
    activity: str
    metrics: Dict[str, Tuple[str, str]]
    warnings: List[str]


FIELDS = [FieldSpec(*x) for x in FIELD_DEFINITIONS]
FIELD_BY_CODE: Dict[str, FieldSpec] = {canonical_code(f.code): f for f in FIELDS if f.kind == "paired" and not f.code.lower().startswith("шкала") and not f.code.lower().startswith("оценка")}
FIELD_BY_DISPLAY = {f.display: f for f in FIELDS}


def extract_docx_content(path: str) -> Tuple[List[str], List[List[List[str]]], str]:
    doc = Document(path)
    paragraphs = [clean_text(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            cells = [clean_text(c.text) for c in r.cells]
            rows.append(cells)
        tables.append(rows)
    full_text = "\n".join(paragraphs)
    for t in tables:
        for r in t:
            full_text += "\n" + " | ".join([c for c in r if c])
    return paragraphs, tables, full_text


def parse_patient_from_text(text: str) -> str:
    m = re.search(r'Ф\.?\s*И\.?\s*О\.?\s*[:\-]\s*(.+)', text, re.I)
    if m:
        return normalize_spaces(m.group(1).splitlines()[0])
    return ""


def parse_datetime_from_text(text: str) -> Optional[dt.datetime]:
    for pattern in DATETIME_PATTERNS:
        m = pattern.search(text)
        if not m:
            continue
        raw = m.group(1).replace(" ", "")
        if "/" in raw:
            dpart, tpart = raw.split("/", 1)
            tpart = tpart.replace(".", ":")
            raw = f"{dpart}/{tpart}"
        try:
            return dt.datetime.strptime(raw, "%d.%m.%Y/%H:%M")
        except ValueError:
            continue
    return None


def parse_date_after_label(text: str, label: str) -> Optional[dt.date]:
    m = re.search(label + r'\s*[:\-]?\s*(\d{2}\.\d{2}\.\d{4})', text, re.I)
    if not m:
        return None
    try:
        return dt.datetime.strptime(m.group(1), "%d.%m.%Y").date()
    except ValueError:
        return None


def detect_year_from_path(path: str) -> str:
    m = re.search(r'ЭТ\s*(20\d{2})', path)
    return m.group(1) if m else ""


def detect_doc_kind(path: str, full_text: str = "") -> str:
    p = normalize_match(path)
    if "2.выпис" in p or "выписн" in p or "повторн" in p or "выписк" in p:
        return "discharge"
    if "1.первич" in p or "первич" in p:
        return "primary"
    t = normalize_match(full_text)
    if "повторная оценка/дата:" in t and re.search(r'повторная оценка/дата:\s*\d{2}\.\d{2}\.\d{4}', t):
        return "discharge"
    return "primary"


def iter_relevant_docx_files(root_dir: str) -> List[str]:
    all_paths = []
    for cur_root, _, files in os.walk(root_dir):
        low = normalize_match(cur_root)
        if "первич" in low or "выпис" in low or "повторн" in low:
            for fn in files:
                if fn.lower().endswith(".docx") and not fn.startswith("~$"):
                    all_paths.append(os.path.join(cur_root, fn))
    return sorted(all_paths)


def detect_activity(full_text: str, paragraphs: List[str]) -> Tuple[str, List[str]]:
    warnings = []
    full_norm = normalize_match(full_text)
    tail_text = "\n".join(paragraphs[-12:]) if paragraphs else full_text
    tail_norm = normalize_match(tail_text)

    if "занятия не показаны" in full_norm:
        return "консультация", warnings
    if "проведена консультация" in full_norm or "консультация по " in tail_norm:
        return "консультация", warnings
    if "индивидуальные занятия с эрготерапевтом" in full_norm or "на индивидуальных занятиях" in full_norm:
        return "индивидуальные", warnings
    if "групповые занятия" in full_norm or "на групповых занятиях" in full_norm:
        return "групповые", warnings

    # "проводились занятия" без уточнения не классифицируем жестко
    if "проводились занятия" in full_norm:
        warnings.append("Проводились занятия, но тип (индивидуальные/групповые) явно не указан")
        return "", warnings

    return "", warnings


def parse_summary_table(table_rows: List[List[str]]) -> Dict[str, Tuple[str, str]]:
    out: Dict[str, Tuple[str, str]] = {}
    if not table_rows:
        return out
    first_row = " | ".join(table_rows[0])
    if "дата" not in normalize_match(first_row):
        return out
    # ожидаем 3 колонки: дата / перв / вып
    for row in table_rows[1:]:
        if not row:
            continue
        label = normalize_match(row[0])
        vals = [parse_value_text(c) for c in row[1:3]]
        v1 = vals[0] or ""
        v2 = vals[1] or ""
        if "sulcs" in label:
            out[FIELD_BY_DISPLAY["Шкала SULCS"].field_id] = (v1, v2)
        elif "fim" in label:
            out[FIELD_BY_DISPLAY["Шкала FIM"].field_id] = (v1, v2)
        elif "eq-5d" in label:
            out[FIELD_BY_DISPLAY["Шкала EQ-5D"].field_id] = (v1, v2)
        elif "выполнение" in label and "удовлетвор" in label:
            out[FIELD_BY_DISPLAY["ОЦЕНКА COPM"].field_id] = (v1, v2)
    return out


def parse_problem_table(table_rows: List[List[str]]) -> Dict[str, Tuple[str, str]]:
    out = {}
    if len(table_rows) < 3:
        return out
    row0 = " | ".join(table_rows[0]).lower()
    row1 = " | ".join(table_rows[1]).lower()
    if "первич" not in row0 and "выполнение 1" not in row1:
        return out

    # первая содержательная строка после шапки
    for row in table_rows[2:]:
        if not row or not clean_text("".join(row)):
            continue
        if "подсчет" in normalize_match(row[0]) or "изменение" in normalize_match(row[0]):
            continue
        v1 = parse_value_text(row[1] if len(row) > 1 else "")
        s1 = parse_value_text(row[2] if len(row) > 2 else "")
        v2 = parse_value_text(row[3] if len(row) > 3 else "")
        s2 = parse_value_text(row[4] if len(row) > 4 else "")
        left = f"{v1}-{s1}" if v1 and s1 else ""
        right = f"{v2}-{s2}" if v2 and s2 else ""
        out[FIELD_BY_DISPLAY["ОЦЕНКА COPM"].field_id] = (left, right)
        break
    return out


def parse_narrative_scale_values(full_text: str, doc_kind: str) -> Dict[str, Tuple[str, str]]:
    out = {}
    text = full_text

    def find_one(patterns):
        for p in patterns:
            m = re.search(p, text, re.I)
            if m:
                return normalize_number_text(m.group(1))
        return ""

    sulcs = find_one([r'SULCS\s*[:\-]\s*(\d+(?:[.,]\d+)?)'])
    fim = find_one([r'FIM[^\d]{0,20}(\d+(?:[.,]\d+)?)'])
    eq = find_one([r'EQ-?5D[^\d]{0,20}(\d+(?:[.,]\d+)?)'])

    if sulcs:
        out[FIELD_BY_DISPLAY["Шкала SULCS"].field_id] = (sulcs, "")
    if fim:
        out[FIELD_BY_DISPLAY["Шкала FIM"].field_id] = (fim, "")
    if eq:
        out[FIELD_BY_DISPLAY["Шкала EQ-5D"].field_id] = (eq, "")
    return out


def parse_mkf_table(table_rows: List[List[str]]) -> Dict[str, Tuple[str, str]]:
    out: Dict[str, Tuple[str, str]] = {}
    if len(table_rows) < 5:
        return out

    # ищем строку с датами
    date_cols = []
    for i, row in enumerate(table_rows[:4]):
        for j, cell in enumerate(row):
            if DATE_RE.fullmatch(clean_text(cell)):
                date_cols.append(j)
    date_cols = sorted(dict.fromkeys(date_cols))
    if not date_cols:
        return out

    # интересны строки, начинающиеся с кода МКФ
    for row in table_rows:
        if not row:
            continue
        raw_code = clean_text(row[0])
        if not raw_code:
            continue
        code_norm = canonical_code(raw_code)
        if not CODE_RE.match(code_norm.replace("е","e")):
            continue
        field = FIELD_BY_CODE.get(code_norm)
        if not field:
            continue

        vin = ""
        vout = ""
        if len(date_cols) >= 1 and date_cols[0] < len(row):
            vin = parse_value_text(row[date_cols[0]]) or ""
        if len(date_cols) >= 2 and date_cols[1] < len(row):
            vout = parse_value_text(row[date_cols[1]]) or ""

        # если есть только одна дата-колонка, это первичный осмотр
        out[field.field_id] = (vin, vout)
    return out


def merge_metric_dict(base: Dict[str, Tuple[str, str]], extra: Dict[str, Tuple[str, str]], prefer_pair: bool = True) -> Dict[str, Tuple[str, str]]:
    res = dict(base)
    for k, (a, b) in extra.items():
        cur = res.get(k, ("", ""))
        # если в extra есть два значения, это почти всегда лучше
        if prefer_pair and a and b:
            res[k] = (a, b)
            continue
        # не затираем хорошие парные значения одиночным числом
        if cur[0] and cur[1] and (not b):
            continue
        res[k] = (
            a if a else cur[0],
            b if b else cur[1],
        )
    return res


def parse_primary_repeat_dates(full_text: str, tables: List[List[List[str]]]) -> Tuple[Optional[dt.date], Optional[dt.date]]:
    p_date = parse_date_after_label(full_text, r'первичн(?:ая|ой)\s+оценка/дата')
    r_date = parse_date_after_label(full_text, r'повторн(?:ая|ой)\s+оценка/дата')

    if p_date or r_date:
        return p_date, r_date

    # fallback по первой таблице с датами
    for t in tables:
        joined = "\n".join(" | ".join(r) for r in t)
        p_date = p_date or parse_date_after_label(joined, r'первичн(?:ая|ой)\s+оценка/дата')
        r_date = r_date or parse_date_after_label(joined, r'повторн(?:ая|ой)\s+оценка/дата')
    return p_date, r_date


def parse_single_doc(path: str) -> ParsedDoc:
    paragraphs, tables, full_text = extract_docx_content(path)
    warnings: List[str] = []

    patient = parse_patient_from_text(full_text)
    if not patient:
        patient = Path(path).stem

    doc_dt = parse_datetime_from_text(full_text)
    exam_date = doc_dt.date() if doc_dt else None
    if doc_dt is None:
        warnings.append("Не удалось извлечь дату/время документа")

    p_eval, r_eval = parse_primary_repeat_dates(full_text, tables)
    activity, aw = detect_activity(full_text, paragraphs)
    warnings.extend(aw)

    metrics: Dict[str, Tuple[str, str]] = {}

    # приоритет 1: итоговая сравнительная таблица с датами
    for t in tables:
        metrics = merge_metric_dict(metrics, parse_summary_table(t), prefer_pair=True)

    # приоритет 2: таблица проблем активности для COPM
    for t in tables:
        metrics = merge_metric_dict(metrics, parse_problem_table(t), prefer_pair=True)

    # приоритет 3: МКФ таблица
    for t in tables:
        metrics = merge_metric_dict(metrics, parse_mkf_table(t), prefer_pair=True)

    # приоритет 4: одиночные шкалы из описательной части, только как fallback
    metrics = merge_metric_dict(metrics, parse_narrative_scale_values(full_text, detect_doc_kind(path, full_text)), prefer_pair=False)

    # заполнить пустые поля
    for f in FIELDS:
        if f.kind != "paired":
            continue
        metrics.setdefault(f.field_id, ("", ""))

    return ParsedDoc(
        path=path,
        year=detect_year_from_path(path),
        doc_kind=detect_doc_kind(path, full_text),
        patient=patient,
        surname=surname_from_fio(patient),
        patient_norm=normalize_patient_name(patient),
        document_dt=doc_dt,
        exam_date=exam_date,
        primary_eval_date=p_eval,
        repeat_eval_date=r_eval,
        activity=activity,
        metrics=metrics,
        warnings=list(dict.fromkeys([w for w in warnings if w])),
    )


def choose_best_doc(docs: List[ParsedDoc]) -> ParsedDoc:
    def sort_key(d: ParsedDoc):
        return (
            1 if d.doc_kind == "discharge" else 0,
            1 if d.repeat_eval_date else 0,
            1 if d.document_dt else 0,
            d.document_dt or dt.datetime.min,
        )
    return sorted(docs, key=sort_key)[-1]


def reduce_docs(all_docs: List[ParsedDoc]) -> Tuple[List[ParsedDoc], List[Tuple[str, str, str]]]:
    """
    Оставляем:
    - все выписки;
    - первичные только если для них не найдено соответствующей выписки.
    Соответствие: тот же пациент, тот же год, и дата первичной оценки в выписке совпадает
    с датой документа первички или отличается не более чем на 7 дней.
    """
    selected: List[ParsedDoc] = []
    warnings = []

    discharges = [d for d in all_docs if d.doc_kind == "discharge"]
    primaries = [d for d in all_docs if d.doc_kind == "primary"]

    # выписки берем после дедупликации по ключу patient+year+repeat/document date
    by_key: Dict[str, List[ParsedDoc]] = {}
    for d in discharges:
        key_date = d.repeat_eval_date or d.exam_date or (d.document_dt.date() if d.document_dt else None)
        key = patient_key_from_name_date(d.year, d.patient_norm, key_date)
        by_key.setdefault(key, []).append(d)

    selected_discharges = []
    for key, docs in by_key.items():
        best = choose_best_doc(docs)
        if len(docs) > 1:
            best.warnings.append(f"Найдено файлов в группе: {len(docs)}. Выбран самый поздний/полный.")
        selected_discharges.append(best)
    selected.extend(selected_discharges)

    # первичные только если не покрыты выпиской
    for p in primaries:
        covered = False
        for d in selected_discharges:
            if p.year != d.year or p.patient_norm != d.patient_norm:
                continue
            if p.exam_date and d.primary_eval_date:
                if abs((p.exam_date - d.primary_eval_date).days) <= 7:
                    covered = True
                    break
            if p.exam_date and d.exam_date:
                if 0 <= (d.exam_date - p.exam_date).days <= 120:
                    covered = True
                    break
        if not covered:
            selected.append(p)

    # финальная дедупликация по patient+year+effective date
    final_map: Dict[str, List[ParsedDoc]] = {}
    for d in selected:
        eff_date = d.repeat_eval_date or d.exam_date or d.primary_eval_date
        key = patient_key_from_name_date(d.year, d.patient_norm, eff_date)
        final_map.setdefault(key, []).append(d)

    final_docs = []
    for key, docs in final_map.items():
        best = choose_best_doc(docs)
        if len(docs) > 1:
            best.warnings.append(f"После объединения осталось дублей: {len(docs)}. Выбран лучший документ.")
        final_docs.append(best)

    final_docs.sort(key=lambda d: (d.year, d.patient_norm, d.exam_date or dt.date.min))
    return final_docs, warnings


def autosize_columns(ws):
    for col_idx, col in enumerate(ws.columns, start=1):
        max_len = 0
        for cell in col:
            val = "" if cell.value is None else str(cell.value)
            max_len = max(max_len, len(val))
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max(max_len + 2, 12), 45)


def style_header(ws):
    fill = PatternFill("solid", fgColor="1F4E78")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def export_xlsx(parsed_docs: List[ParsedDoc], out_path: str):
    wb = Workbook()
    ws = wb.active
    ws.title = "cases_wide"

    paired_fields = [f for f in FIELDS if f.kind == "paired"]
    headers = [
        "Год", "Тип документа", "Дата осмотра", "Дата/время документа",
        "Первичная оценка", "Повторная оценка",
        "Пациент", "Фамилия", "Нормализованный пациент", "Ключ сопоставления",
        "Характер занятий", "Источник файла", "Предупреждения"
    ]
    for f in paired_fields:
        headers.append(f"{f.display} | поступление")
        headers.append(f"{f.display} | выписка")
    ws.append(headers)

    for d in parsed_docs:
        eff_date = d.repeat_eval_date or d.exam_date or d.primary_eval_date
        row = [
            d.year,
            "выписка" if d.doc_kind == "discharge" else "первичный",
            eff_date.isoformat() if eff_date else "",
            d.document_dt.strftime("%d.%m.%Y %H:%M") if d.document_dt else "",
            d.primary_eval_date.isoformat() if d.primary_eval_date else "",
            d.repeat_eval_date.isoformat() if d.repeat_eval_date else "",
            d.patient,
            d.surname,
            d.patient_norm,
            patient_key_from_name_date(d.year, d.patient_norm, eff_date),
            d.activity,
            d.path,
            " | ".join(d.warnings),
        ]
        for f in paired_fields:
            a, b = d.metrics.get(f.field_id, ("", ""))
            row.extend([a, b])
        ws.append(row)

    style_header(ws)
    autosize_columns(ws)
    ws.freeze_panes = "A2"

    # raw
    ws2 = wb.create_sheet("files_raw")
    ws2.append(["Год", "Тип документа", "Пациент", "Дата/время документа", "Дата документа", "Первичная оценка", "Повторная оценка", "Характер занятий", "Файл", "Предупреждения"])
    for d in parsed_docs:
        ws2.append([
            d.year, d.doc_kind, d.patient,
            d.document_dt.strftime("%d.%m.%Y %H:%M") if d.document_dt else "",
            d.exam_date.isoformat() if d.exam_date else "",
            d.primary_eval_date.isoformat() if d.primary_eval_date else "",
            d.repeat_eval_date.isoformat() if d.repeat_eval_date else "",
            d.activity, d.path, " | ".join(d.warnings)
        ])
    style_header(ws2)
    autosize_columns(ws2)
    ws2.freeze_panes = "A2"

    ws3 = wb.create_sheet("warnings")
    ws3.append(["Тип", "Год", "Пациент", "Файл", "Сообщение"])
    for d in parsed_docs:
        for w in d.warnings:
            ws3.append(["document_warning", d.year, d.patient, d.path, w])
    style_header(ws3)
    autosize_columns(ws3)
    ws3.freeze_panes = "A2"

    wb.save(out_path)


def build_arg_parser():
    p = argparse.ArgumentParser()
    p.add_argument("--root", default=DEFAULT_ROOT_DIR)
    p.add_argument("--output", default=DEFAULT_OUTPUT_XLSX)
    return p


def main():
    args = build_arg_parser().parse_args()
    root = args.root
    files = iter_relevant_docx_files(root)
    if not files:
        print(f"Не найдено .docx в {root}")
        return 1

    parsed = []
    for i, path in enumerate(files, start=1):
        try:
            parsed.append(parse_single_doc(path))
        except Exception as e:
            year = detect_year_from_path(path)
            parsed.append(ParsedDoc(
                path=path, year=year, doc_kind=detect_doc_kind(path, full_text), patient=Path(path).stem,
                surname=surname_from_fio(Path(path).stem), patient_norm=normalize_patient_name(Path(path).stem),
                document_dt=None, exam_date=None, primary_eval_date=None, repeat_eval_date=None,
                activity="", metrics={f.field_id: ("", "") for f in FIELDS if f.kind=="paired"},
                warnings=[f"Ошибка разбора документа: {e}"]
            ))

    selected, _ = reduce_docs(parsed)
    export_xlsx(selected, args.output)
    print(f"Готово: {args.output}")
    print(f"Обработано файлов: {len(files)}")
    print(f"Итоговых случаев: {len(selected)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
